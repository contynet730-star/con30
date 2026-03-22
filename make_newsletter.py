from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── ページ設定（A4、余白狭め）──────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.top_margin    = Cm(1.0)
section.bottom_margin = Cm(1.0)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)

# ── ヘルパー ─────────────────────────────────────────
def set_font(run, size, bold=False, color=None, name="MS Gothic"):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

def para_space(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, val in [("top",top),("bottom",bottom),("left",left),("right",right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val.get("val","single"))
            el.set(qn("w:sz"),    val.get("sz","4"))
            el.set(qn("w:color"), val.get("color","000000"))
            borders.append(el)
    tcPr.append(borders)

def add_run_in_para(p, text, size, bold=False, color=None, name="MS Gothic"):
    r = p.add_run(text)
    set_font(r, size, bold, color, name)
    return r

# ══════════════════════════════════════════════════════════
# ① ヘッダー帯「全教職員向け…」
# ══════════════════════════════════════════════════════════
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
cell = tbl.rows[0].cells[0]
set_cell_bg(cell, "2E4057")  # 濃紺
cell.width = Cm(18)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(p, before=1, after=1, line=12)
r = p.add_run("全教職員向け「生活指導だより」")
set_font(r, 10, bold=True, color=(255,255,255))

# ══════════════════════════════════════════════════════════
# ② タイトル帯（緑黒板風）
# ══════════════════════════════════════════════════════════
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = "Table Grid"

# 左：タイトル
c_left = tbl2.rows[0].cells[0]
c_left.width = Cm(13)
set_cell_bg(c_left, "1B4332")  # 深緑
p = c_left.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(p, before=2, after=0, line=14)
r = p.add_run("「あん・しん」「す・て・き」な\n")
set_font(r, 13, bold=True, color=(255,230,50))
r2 = p.add_run("学校づくり")
set_font(r2, 15, bold=True, color=(255,230,50))

p2 = c_left.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(p2, before=1, after=2, line=10)
r3 = p2.add_run("あん…安心　しん…信頼　す…素早く　て…丁寧に　き…協動的")
set_font(r3, 7.5, color=(200,230,200))

# 右：号数・日付・発行元
c_right = tbl2.rows[0].cells[1]
c_right.width = Cm(5)
set_cell_bg(c_right, "1B4332")
p3 = c_right.paragraphs[0]
p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
para_space(p3, before=2, after=0, line=13)
r4 = p3.add_run("第４号\n令和８年３月\n練馬区教育委員会\n生活指導担当指導主事")
set_font(r4, 8, color=(255,255,255))

# ══════════════════════════════════════════════════════════
# ③ リード文
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
para_space(p, before=3, after=1, line=11)
r = p.add_run(
    "　このたび区教育委員会より「練馬区立学校における学校の決まりや校則等の見直し・確認について"
    "（通知）」（令和８年３月12日　7練教教指第5589号）が発出されました。文部科学省・生徒指導提要の"
    "考え方と合わせ、各校での取組にご活用ください。"
)
set_font(r, 7.5)

# ══════════════════════════════════════════════════════════
# ④ 2カラム構成（左：区通知５観点　右：要注意校則例）
# ══════════════════════════════════════════════════════════
tbl3 = doc.add_table(rows=1, cols=2)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.style = "Table Grid"

# ─── 左セル ───────────────────────────────────────────
cl = tbl3.rows[0].cells[0]
cl.width = Cm(9)
set_cell_bg(cl, "EAF4FB")

# 左タイトル
p = cl.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(p, before=2, after=1, line=11)
r = p.add_run("１．区通知が示す５つの観点")
set_font(r, 8.5, bold=True, color=(0,70,140))

items = [
    ("①", "校則は教育目的を実現する過程で定める学習上・生活上の規律である"),
    ("②", "児童生徒の人権および意見を尊重したものである"),
    ("③", "児童会・生徒会等を通じた意見収集を行い、広く関係者の理解を得る"),
    ("④", "昨年度見直したとしても、現在の実態に応じて改めて見直しを図る"),
    ("⑤", "意義や適切な説明ができない校則は再度検討する"),
]
for num, text in items:
    p = cl.add_paragraph()
    para_space(p, before=0, after=1, line=11)
    r1 = p.add_run(f"{num} ")
    set_font(r1, 8, bold=True, color=(0,100,180))
    r2 = p.add_run(text)
    set_font(r2, 7.5)

p = cl.add_paragraph()
para_space(p, before=2, after=1, line=11)
r = p.add_run("２．生徒指導提要・文科省の考え方")
set_font(r, 8.5, bold=True, color=(0,70,140))

points = [
    "★ 校則は「教育目的実現のための規律」。\n　児童生徒が意義を理解し自律的に行動できる力を育てることが目標。",
    "★ 合理的な説明ができない校則は積極的に見直す（文科省 令和３年）",
    "★ 見直しに生徒・保護者が参加することで規範意識と自治的能力が育まれる",
]
for pt in points:
    p = cl.add_paragraph()
    para_space(p, before=1, after=1, line=11)
    r = p.add_run(pt)
    set_font(r, 7.5)

# ─── 右セル ───────────────────────────────────────────
cr = tbl3.rows[0].cells[1]
cr.width = Cm(9)
set_cell_bg(cr, "FFF8E7")

p = cr.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(p, before=2, after=1, line=11)
r = p.add_run("２．こんな校則・きまりは要注意")
set_font(r, 8.5, bold=True, color=(180,60,0))

p = cr.add_paragraph()
para_space(p, before=0, after=1, line=11)
r = p.add_run("区別紙では、合理的な根拠や説明が示されていないことから、"
              "管理統制と受け止められたり誤解につながる可能性があるものとして以下を例示。")
set_font(r, 7.5)

caution = [
    "□ 授業中の行動（水分補給・タブレット使用等）にも教師の許可が必要",
    "□ ゴムやヘアピンの色を「黒・紺・茶」に限定するなど詳細な身だしなみ規定",
    "□ 「中学生らしい服装」など抽象的な表現による指定",
    "□ 清掃中・給食準備中の私語を禁止し「無言」を義務づける",
    "□ 他学級・他学年フロアへの入室を禁止する表現",
]
for c in caution:
    p = cr.add_paragraph()
    para_space(p, before=0, after=1, line=11)
    r = p.add_run(c)
    set_font(r, 7.5, color=(140,30,0))

p = cr.add_paragraph()
para_space(p, before=2, after=1, line=11)
r = p.add_run("【ポイント】")
set_font(r, 8, bold=True, color=(180,60,0))
p2 = cr.add_paragraph()
para_space(p2, before=0, after=1, line=11)
r2 = p2.add_run("教師がその背景・理由を理解し、児童生徒に意義を"
                "説明できること、適宜見直す機会が設けられていることが必要です。")
set_font(r2, 7.5)

# ══════════════════════════════════════════════════════════
# ⑤ 見直しステップ（横並び5ステップ）
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
para_space(p, before=3, after=1, line=11)
r = p.add_run("３．校則見直しの実践ステップ")
set_font(r, 8.5, bold=True, color=(0,70,140))

steps = [
    ("STEP1\n現状把握", "校則を一覧化し\n根拠・目的を確認\n説明できないものを点検"),
    ("STEP2\n意見収集", "生徒会・保護者\nアンケート等で\n広く意見を集める"),
    ("STEP3\n協議・検討", "管理職・生徒代表\nが参加し多角的に\n人権的観点で協議"),
    ("STEP4\n改定・周知", "HP等で公開し\n全保護者・生徒に\nプロセスも説明"),
    ("STEP5\n継続見直し", "毎年度機会を確保\n「昨年やった」で\n終わらせない"),
]

tbl4 = doc.add_table(rows=1, cols=5)
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl4.style = "Table Grid"
colors = ["1A5276","1F618D","2471A3","2980B9","3498DB"]
for i, (title, body) in enumerate(steps):
    c = tbl4.rows[0].cells[i]
    set_cell_bg(c, colors[i])
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(p, before=2, after=0, line=11)
    r1 = p.add_run(title + "\n")
    set_font(r1, 7.5, bold=True, color=(255,230,100))
    r2 = p.add_run(body)
    set_font(r2, 7, color=(255,255,255))
    para_space(p, before=2, after=2, line=11)

# ══════════════════════════════════════════════════════════
# ⑥ 生徒主体の見直し手続き
# ══════════════════════════════════════════════════════════
tbl5 = doc.add_table(rows=1, cols=1)
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl5.style = "Table Grid"
cb = tbl5.rows[0].cells[0]
set_cell_bg(cb, "F0FFF0")
p = cb.paragraphs[0]
para_space(p, before=2, after=1, line=11)
r = p.add_run("【生徒主体の見直し手続き例（区別紙より）】　")
set_font(r, 8, bold=True, color=(0,100,0))
r2 = p.add_run("学級・学年での話し合い　→　児童会・生徒会　→　代表委員会（中央委員会）　→　校長へ要望　→　検討・改定")
set_font(r2, 7.5, color=(0,80,0))
p2 = cb.add_paragraph()
para_space(p2, before=0, after=2, line=11)
r3 = p2.add_run("　少数意見も尊重し、公正な手続きの中でみんなでより良い学校生活をつくることが大切です。")
set_font(r3, 7.5)

# ══════════════════════════════════════════════════════════
# ⑦ 区への相談案件＋管理職へ一報
# ══════════════════════════════════════════════════════════
tbl6 = doc.add_table(rows=1, cols=2)
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl6.style = "Table Grid"

cl2 = tbl6.rows[0].cells[0]
cl2.width = Cm(11)
set_cell_bg(cl2, "F5EEF8")

p = cl2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(p, before=2, after=1, line=11)
r = p.add_run("★ 区教委に寄せられた最近の生活指導案件（校則関係）")
set_font(r, 8, bold=True, color=(80,0,120))

cases = [
    ("☑ 「なぜその校則があるのか説明してもらえなかった」（保護者）",
     "→ まず教員間で説明できるか確認。できない場合は見直しの対象（通知 観点⑤）"),
    ("☑ 「ゴムの色まで指定されるのは管理しすぎ」（生徒・保護者）",
     "→ 区別紙の例示事案。根拠・教育的意義を確認し、説明できなければ見直しへ"),
]
for q, a in cases:
    p = cl2.add_paragraph()
    para_space(p, before=1, after=0, line=11)
    r1 = p.add_run(q)
    set_font(r1, 7.5, bold=True)
    p2 = cl2.add_paragraph()
    para_space(p2, before=0, after=2, line=11)
    r2 = p2.add_run(a)
    set_font(r2, 7.5, color=(80,0,120))

cr2 = tbl6.rows[0].cells[1]
cr2.width = Cm(7)
set_cell_bg(cr2, "FDEDEC")

p = cr2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(p, before=2, after=1, line=11)
r = p.add_run("⚠ 管理職へ一報を")
set_font(r, 8.5, bold=True, color=(180,0,0))

msgs = [
    "校則に関するトラブルは担任だけで抱え込まず、学年主任・主任教諭・管理職へ速やかに報告",
    "ホウレンソウ（報告・連絡・相談）を徹底し、組織として対応方針を共有する",
    "区通知に基づき、見直し検討につなげる",
]
for m in msgs:
    p = cr2.add_paragraph()
    para_space(p, before=1, after=1, line=11)
    r = p.add_run("● " + m)
    set_font(r, 7.5)

# ══════════════════════════════════════════════════════════
# ⑧ フッター注記
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
para_space(p, before=2, after=0, line=10)
r = p.add_run(
    "※ 生徒指導提要（令和４年12月改訂版）は文部科学省HPよりダウンロードできます。"
    "　※ 区通知（7練教教指第5589号）の別紙も合わせて各校でご確認ください。"
    "　※ ご不明点は区教育委員会生活指導担当指導主事までお問い合わせください。"
)
set_font(r, 6.5, color=(80,80,80))

doc.save("/home/user/con30/生活指導だより第4号_3月号.docx")
print("Done")
